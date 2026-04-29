from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        if part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            continue
        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            continue
        paragraph.add_run(part)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | tuple[None, int]:
    table_lines: list[str] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].rstrip())
        idx += 1

    if len(table_lines) < 2:
        return None, start

    separator = table_lines[1].replace("|", "").replace("-", "").replace(":", "").strip()
    if separator:
        return None, start

    rows: list[list[str]] = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    return rows, idx


def add_table(document: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    body = rows[2:]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for col, value in enumerate(headers):
        add_inline_runs(table.rows[0].cells[col].paragraphs[0], value)

    for row in body:
        tr = table.add_row().cells
        for col, value in enumerate(row):
            add_inline_runs(tr[col].paragraphs[0], value)


def try_add_figure(document: Document, markdown_line: str, base_dir: Path) -> None:
    image_paths = re.findall(r"`([^`]+\.(?:png|jpg|jpeg))`", markdown_line, flags=re.IGNORECASE)
    for relative in image_paths:
        candidate = (base_dir / relative).resolve()
        if candidate.exists():
            try:
                document.add_picture(str(candidate), width=Cm(16))
            except Exception:
                note = document.add_paragraph(style="Intense Quote")
                note.add_run(f"Immagine non inserita automaticamente: {candidate}")


def convert_markdown(input_path: Path, output_path: Path) -> None:
    text = input_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Cm(1.8))

    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            table_rows, next_idx = parse_table(lines, idx)
            if table_rows:
                add_table(document, table_rows)
                idx = next_idx
                continue

        if stripped == "---":
            document.add_paragraph()
            idx += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            paragraph = document.add_paragraph(style=f"Heading {min(level, 9)}")
            add_inline_runs(paragraph, stripped[level:].strip())
            idx += 1
            continue

        if stripped.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:].strip())
            try_add_figure(document, stripped, input_path.parent)
            idx += 1
            continue

        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            candidate = lines[idx].strip()
            if not candidate:
                break
            if candidate.startswith(("#", "- ", "* ", "|")) or candidate == "---":
                break
            paragraph_lines.append(candidate)
            idx += 1

        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, " ".join(paragraph_lines))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert a simple Markdown file to DOCX.")
    parser.add_argument("input_path")
    parser.add_argument("output_path")
    args = parser.parse_args()

    convert_markdown(Path(args.input_path), Path(args.output_path))


if __name__ == "__main__":
    main()
